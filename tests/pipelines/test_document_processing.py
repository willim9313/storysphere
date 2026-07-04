"""Unit tests for the document processing pipeline.

Tests cover:
- loader helpers (mocked file I/O)
- chapter_detector heuristics
- chunker split/merge logic
- DocumentProcessingPipeline with a mock loader
"""

from __future__ import annotations
from unittest.mock import MagicMock, patch

import pytest

from storysphere.domain.documents import ChapterRole
from storysphere.pipelines.document_processing.chapter_detector import (
    ChapterSpan,
    _is_inline_title,
    _match_heading,
    detect_chapters,
)
from storysphere.pipelines.document_processing.chunker import chunk_segments


# ── chapter_detector ────────────────────────────────────────────────────────


class TestDetectChapters:
    def test_empty_input_returns_empty(self):
        assert detect_chapters([]) == []

    def test_no_headings_produces_single_chapter(self):
        segments = [(0, "Once upon a time."), (1, "The hero walked.")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 1
        assert chapters[0].chapter_number == 1
        assert chapters[0].title is None
        assert len(chapters[0].segments) == 2

    def test_chapter_heading_splits_correctly(self):
        segments = [
            (0, "Chapter 1"),
            (1, "The beginning."),
            (2, "Chapter 2"),
            (3, "The middle."),
        ]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2
        assert chapters[0].chapter_number == 1
        assert len(chapters[0].segments) == 1
        assert chapters[1].chapter_number == 2

    def test_cjk_chapter_heading(self):
        segments = [(0, "第一章"), (1, "龍出現了。"), (2, "第二章"), (3, "戰鬥開始。")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2

    def test_roman_numeral_heading(self):
        segments = [(0, "I"), (1, "Text A."), (2, "II"), (3, "Text B.")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2

    def test_prologue_heading(self):
        segments = [(0, "Prologue"), (1, "Before it all."), (2, "Chapter 1"), (3, "It started.")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2
        assert chapters[0].chapter_number == 1

    def test_content_before_first_heading(self):
        """Text before any chapter heading becomes chapter 1."""
        segments = [(0, "Author's note."), (1, "Chapter 1"), (2, "Story starts.")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2
        assert chapters[0].segments[0][1] == "Author's note."

    # ── new patterns (old_version parity) ──

    def test_volume_book_part_headings(self):
        """Volume/Book/Part/Act/Scene as standalone headings."""
        for heading in ["Volume 1", "Book 3", "Part 2", "Act 1", "Scene 5", "Vol II"]:
            segments = [(0, heading), (1, "Content.")]
            chapters = detect_chapters(segments)
            assert len(chapters) == 1, f"Failed for: {heading}"

    def test_chapter_with_title_separator(self):
        """Chapter heading followed by separator and title."""
        for heading in [
            "Chapter 3 - The Return",
            "Chapter 5: Awakening",
            "Ch. 2: Title",
            "Chapter IV: The Quest",
        ]:
            segments = [(0, heading), (1, "Content.")]
            chapters = detect_chapters(segments)
            assert len(chapters) == 1, f"Failed for: {heading}"

    def test_cjk_chapter_with_title(self):
        """Chinese chapter heading with separator and title."""
        for heading in ["第五章：歸來", "第1章: 開始", "第三節-啟程"]:
            segments = [(0, heading), (1, "內容。")]
            chapters = detect_chapters(segments)
            assert len(chapters) == 1, f"Failed for: {heading}"

    def test_compound_volume_chapter(self):
        """Compound volume + chapter format."""
        for heading in [
            "Volume 1, Chapter 5",
            "Book 2, Act 3",
            "Part 1 Chapter 1",
        ]:
            segments = [(0, heading), (1, "Content.")]
            chapters = detect_chapters(segments)
            assert len(chapters) == 1, f"Failed for: {heading}"

    def test_extended_cjk_characters(self):
        """Chinese chapter with 萬/零/〇 characters."""
        for heading in ["第零章", "第〇章"]:
            segments = [(0, heading), (1, "內容。")]
            chapters = detect_chapters(segments)
            assert len(chapters) == 1, f"Failed for: {heading}"

    # ── styled_heading_indices (DOCX "Heading N" style priority) ──

    def test_styled_heading_with_no_regex_match_becomes_titled_chapter(self):
        """A styled heading whose text matches no regex (e.g. '楔子') still
        starts a new chapter, using the whole line as the title."""
        segments = [(0, "楔子"), (1, "很久很久以前。"), (2, "正文"), (3, "故事開始了。")]
        chapters = detect_chapters(segments, styled_heading_indices={0})
        assert len(chapters) == 1
        assert chapters[0].title == "楔子"
        assert chapters[0].segments == [(1, "很久很久以前。"), (2, "正文"), (3, "故事開始了。")]

    def test_styled_heading_overrides_regex_for_untitled_body_line(self):
        """A plain body line with no chapter-heading shape still becomes a
        chapter boundary when the source format flags it as a heading style."""
        segments = [(0, "Some Custom Title"), (1, "Body text.")]
        # Without the style hint this line matches no heading pattern at all.
        assert len(detect_chapters(segments)) == 1
        assert detect_chapters(segments)[0].title is None

        chapters = detect_chapters(segments, styled_heading_indices={0})
        assert len(chapters) == 1
        assert chapters[0].title == "Some Custom Title"

    def test_styled_heading_with_inline_separator_title_still_extracted(self):
        """Regex-extractable titles still win over the raw-line fallback."""
        segments = [(0, "Chapter 3 - The Return"), (1, "Content.")]
        chapters = detect_chapters(segments, styled_heading_indices={0})
        assert len(chapters) == 1
        assert chapters[0].title == "The Return"

    def test_styled_heading_bare_number_still_peeks_next_line_for_title(self):
        """A styled heading that IS a bare chapter-number label keeps the
        existing inline-title-peek behavior instead of using itself as title."""
        segments = [(0, "Chapter 1"), (1, "The Awakening"), (2, "Body text.")]
        chapters = detect_chapters(segments, styled_heading_indices={0})
        assert len(chapters) == 1
        assert chapters[0].title == "The Awakening"
        assert chapters[0].segments == [(2, "Body text.")]

    def test_no_styled_heading_indices_matches_default_behavior(self):
        """Passing None/empty behaves identically to omitting the argument."""
        segments = [(0, "Chapter 1"), (1, "Text.")]
        assert detect_chapters(segments) == detect_chapters(segments, styled_heading_indices=None)
        assert detect_chapters(segments) == detect_chapters(segments, styled_heading_indices=set())

    # ── chapter-level role classification (toc/preface/afterword) ──

    def test_toc_preface_body_afterword_classified_as_separate_chapters(self):
        segments = [
            (0, "目錄"),
            (1, "這裡條列本書所有章節標題，這裡條列本書所有章節標題。"),
            (2, "推薦序：一位讀者的告白"),
            (3, "這本書非常精彩值得一讀值得一讀值得一讀值得一讀值得一讀。"),
            (4, "第一章 開始"),
            (5, "故事開始了故事開始了故事開始了故事開始了故事開始了。"),
            (6, "後記"),
            (7, "感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持。"),
        ]
        chapters = detect_chapters(segments)
        assert [c.role for c in chapters] == [
            ChapterRole.toc,
            ChapterRole.preface,
            ChapterRole.body,
            ChapterRole.afterword,
        ]

    def test_prologue_and_epilogue_stay_body_role(self):
        segments = [
            (0, "Prologue"),
            (1, "It all began on a stormy night many years ago."),
            (2, "Chapter 1"),
            (3, "The story begins in earnest here today."),
        ]
        chapters = detect_chapters(segments)
        assert [c.role for c in chapters] == [ChapterRole.body, ChapterRole.body]

    def test_default_chapter_has_body_role(self):
        segments = [(0, "Once upon a time."), (1, "The hero walked.")]
        chapters = detect_chapters(segments)
        assert chapters[0].role == ChapterRole.body


# ── _match_heading title extraction ─────────────────────────────────────────


class TestMatchHeadingTitleExtraction:
    def test_en_chapter_with_separator_extracts_title(self):
        assert _match_heading("Chapter 3 - The Return") == (True, "The Return", ChapterRole.body)
        assert _match_heading("Ch. 2: Awakening") == (True, "Awakening", ChapterRole.body)
        assert _match_heading("Chapter IV: The Quest") == (True, "The Quest", ChapterRole.body)

    def test_cjk_chapter_with_separator_extracts_title(self):
        assert _match_heading("第五章：歸來") == (True, "歸來", ChapterRole.body)
        assert _match_heading("第1章: 開始") == (True, "開始", ChapterRole.body)
        assert _match_heading("第三節-啟程") == (True, "啟程", ChapterRole.body)

    def test_cjk_chapter_with_space_extracts_title(self):
        assert _match_heading("第二章 茅草、木頭與一封奇信") == (
            True, "茅草、木頭與一封奇信", ChapterRole.body,
        )
        assert _match_heading("第三章 英雄登場") == (True, "英雄登場", ChapterRole.body)

    def test_number_only_headings_have_no_title(self):
        assert _match_heading("Chapter 1") == (True, None, ChapterRole.body)
        assert _match_heading("第一章") == (True, None, ChapterRole.body)
        assert _match_heading("Prologue") == (True, None, ChapterRole.body)
        assert _match_heading("Volume 1") == (True, None, ChapterRole.body)

    def test_non_heading_returns_false(self):
        is_h, title, role = _match_heading("Some random body text that is long enough.")
        assert is_h is False
        assert title is None
        assert role == ChapterRole.body

    def test_over_80_chars_not_heading(self):
        long = "第" + "一" * 81 + "章 title"
        assert _match_heading(long) == (False, None, ChapterRole.body)

    def test_toc_and_front_back_matter_markers_classified(self):
        assert _match_heading("目錄") == (True, "目錄", ChapterRole.toc)
        assert _match_heading("Table of Contents") == (True, "Table of Contents", ChapterRole.toc)
        assert _match_heading("序") == (True, "序", ChapterRole.preface)
        assert _match_heading("推薦序：一位讀者的告白") == (
            True, "推薦序：一位讀者的告白", ChapterRole.preface,
        )
        assert _match_heading("Foreword") == (True, "Foreword", ChapterRole.preface)
        assert _match_heading("後記") == (True, "後記", ChapterRole.afterword)
        assert _match_heading("Afterword") == (True, "Afterword", ChapterRole.afterword)

    def test_prologue_epilogue_kept_as_body_not_preface(self):
        """Prologue/epilogue are narrative content, unlike preface/afterword."""
        assert _match_heading("Prologue") == (True, None, ChapterRole.body)
        assert _match_heading("Epilogue") == (True, None, ChapterRole.body)


# ── _is_inline_title heuristic ───────────────────────────────────────────────


class TestIsInlineTitle:
    def test_short_cjk_no_punct_is_title(self):
        assert _is_inline_title("茅草、木頭與一封奇信") is True
        assert _is_inline_title("英雄登場") is True
        assert _is_inline_title("The Return") is True

    def test_ends_with_sentence_punct_not_title(self):
        assert _is_inline_title("龍出現了。") is False
        assert _is_inline_title("The beginning.") is False
        assert _is_inline_title("戰鬥開始！") is False

    def test_too_long_not_title(self):
        assert _is_inline_title("茅" * 31) is False

    def test_empty_not_title(self):
        assert _is_inline_title("") is False
        assert _is_inline_title("   ") is False

    def test_starts_with_non_word_not_title(self):
        assert _is_inline_title("— 47 —") is False
        assert _is_inline_title("***") is False


# ── detect_chapters inline title promotion ───────────────────────────────────


class TestDetectChaptersInlineTitlePromotion:
    def test_short_title_segment_after_heading_promoted(self):
        """First body segment promoted to title when heading has no title."""
        segments = [(0, "第二章"), (1, "茅草與奇信"), (2, "三兄弟離家的那年。")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 1
        assert chapters[0].title == "茅草與奇信"
        # title segment should NOT appear in body
        body_texts = [t for _, t in chapters[0].segments]
        assert "茅草與奇信" not in body_texts

    def test_sentence_segment_after_heading_stays_as_body(self):
        """First segment ending with 。 is NOT promoted — stays as body content."""
        segments = [(0, "第一章"), (1, "龍出現了。"), (2, "繼續戰鬥。")]
        chapters = detect_chapters(segments)
        assert chapters[0].title is None
        assert any("龍出現了" in t for _, t in chapters[0].segments)

    def test_inline_title_extracted_from_heading_line(self):
        """Title embedded in heading line (第X章 title) is extracted directly."""
        segments = [(0, "第二章 茅草、木頭與一封奇信"), (1, "三兄弟離家的那年。")]
        chapters = detect_chapters(segments)
        assert len(chapters) == 1
        assert chapters[0].title == "茅草、木頭與一封奇信"

    def test_inline_title_not_promoted_if_heading_already_has_title(self):
        """When heading itself has a title, next segment is always body."""
        segments = [(0, "第五章：歸來"), (1, "英雄"), (2, "故事繼續。")]
        chapters = detect_chapters(segments)
        assert chapters[0].title == "歸來"
        # "英雄" should remain in segments (heading already had a title)
        body_texts = [t for _, t in chapters[0].segments]
        assert "英雄" in body_texts

    def test_multiple_chapters_each_get_inline_title(self):
        segments = [
            (0, "第一章"), (1, "英雄登場"), (2, "内容一。"),
            (3, "第二章"), (4, "反派出現"), (5, "内容二。"),
        ]
        chapters = detect_chapters(segments)
        assert len(chapters) == 2
        assert chapters[0].title == "英雄登場"
        assert chapters[1].title == "反派出現"


# ── chunker ─────────────────────────────────────────────────────────────────


class TestChunkSegments:
    def test_empty_segments_return_empty(self):
        result = chunk_segments([], chapter_number=1)
        assert result == []

    def test_normal_segment_becomes_paragraph(self):
        segments = [(0, "This is a longer sentence about something interesting enough to keep.")]
        paragraphs = chunk_segments(segments, chapter_number=1)
        assert len(paragraphs) == 1
        assert paragraphs[0].chapter_number == 1
        assert paragraphs[0].position == 0

    def test_short_segments_are_merged(self):
        """Two short segments should be merged into one paragraph."""
        segments = [(0, "Hello world."), (1, "Goodbye world.")]
        paragraphs = chunk_segments(segments, chapter_number=2, min_chars=10)
        # "Hello world." + " " + "Goodbye world." = 27 chars → above 10, merged into 1
        assert len(paragraphs) == 1
        assert "Hello world." in paragraphs[0].text

    def test_long_segment_is_split(self):
        """A segment exceeding max_chars should be split into multiple paragraphs."""
        # Build a long text > 100 chars with sentence boundaries
        long_text = "Sentence one. Sentence two. Sentence three. " * 5
        segments = [(0, long_text)]
        paragraphs = chunk_segments(segments, chapter_number=3, max_chars=100)
        assert len(paragraphs) > 1

    def test_positions_are_sequential(self):
        segments = [(i, f"This is paragraph number {i} with enough characters here.") for i in range(5)]
        paragraphs = chunk_segments(segments, chapter_number=1)
        positions = [p.position for p in paragraphs]
        assert positions == list(range(len(paragraphs)))

    def test_whitespace_only_segments_are_skipped(self):
        segments = [(0, "   "), (1, "Real content here."), (2, "\n\t")]
        paragraphs = chunk_segments(segments, chapter_number=1)
        assert all(p.text.strip() for p in paragraphs)


# ── DocumentProcessingPipeline (integration-lite) ────────────────────────────


class TestDocumentProcessingPipeline:
    @pytest.mark.asyncio
    async def test_unsupported_extension_raises(self, tmp_path):
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        bad_file = tmp_path / "novel.xyz"
        bad_file.write_text("hello")

        pipeline = DocumentProcessingPipeline()
        with pytest.raises(ValueError, match="Unsupported file type"):
            await pipeline.run(bad_file)

    @pytest.mark.asyncio
    async def test_missing_file_raises(self, tmp_path):
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        pipeline = DocumentProcessingPipeline()
        with pytest.raises(FileNotFoundError):
            await pipeline.run(tmp_path / "missing.pdf")

    @pytest.mark.asyncio
    async def test_chapter_role_propagates_end_to_end(self, tmp_path):
        """A real .txt file with toc/preface/body/afterword sections ends up
        with the matching ``Chapter.role`` on the final Document."""
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        txt_file = tmp_path / "novel.txt"
        txt_file.write_text(
            "目錄\n"
            "這裡條列本書所有章節標題，這裡條列本書所有章節標題，這裡條列本書所有章節標題，這裡條列本書所有章節標題。\n"
            "推薦序：一位讀者的告白\n"
            "這本書非常精彩值得一讀值得一讀值得一讀值得一讀值得一讀值得一讀值得一讀值得一讀值得一讀值得一讀。\n"
            "第一章 開始\n"
            "故事開始了故事開始了故事開始了故事開始了故事開始了故事開始了故事開始了故事開始了故事開始了故事開始了。\n"
            "後記\n"
            "感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持感謝各位讀者的支持。\n",
            encoding="utf-8",
        )

        pipeline = DocumentProcessingPipeline()
        doc = await pipeline.run(txt_file)

        assert [c.role for c in doc.chapters] == [
            ChapterRole.toc,
            ChapterRole.preface,
            ChapterRole.body,
            ChapterRole.afterword,
        ]

    @pytest.mark.asyncio
    async def test_docx_processing(self, tmp_path):
        """Create a minimal DOCX and verify it produces a Document."""
        try:
            import docx  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        import docx as docx_lib

        doc_file = tmp_path / "test.docx"
        wdoc = docx_lib.Document()
        wdoc.add_paragraph("Chapter 1")
        wdoc.add_paragraph("The hero woke up on a cold morning and looked around the room.")
        wdoc.add_paragraph("Chapter 2")
        wdoc.add_paragraph("The journey began at dawn. The path was steep and treacherous.")
        wdoc.save(str(doc_file))

        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        pipeline = DocumentProcessingPipeline()
        result = await pipeline.run(doc_file)

        assert result.title == "test"
        assert result.total_chapters >= 1
        assert result.total_paragraphs >= 1


# ── loader: DocumentMeta extraction ─────────────────────────────────────────


class TestLoaderMeta:
    def test_load_pdf_extracts_author_and_title(self, tmp_path):
        """load_pdf reads /Author and /Title from PDF metadata."""
        from storysphere.pipelines.document_processing.loader import load_pdf

        fake_pdf = tmp_path / "novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Chapter 1\nSome content here."

        mock_reader = MagicMock()
        mock_reader.metadata = {"/Author": "Jane Austen", "/Title": "Pride and Prejudice"}
        mock_reader.pages = [mock_page]

        with patch("pypdf.PdfReader", return_value=mock_reader):
            segments, meta = load_pdf(fake_pdf)

        assert meta.author == "Jane Austen"
        assert meta.title == "Pride and Prejudice"
        assert len(segments) > 0

    def test_load_pdf_handles_missing_metadata(self, tmp_path):
        """load_pdf returns None fields when PDF has no metadata."""
        from storysphere.pipelines.document_processing.loader import load_pdf

        fake_pdf = tmp_path / "novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Some content."

        mock_reader = MagicMock()
        mock_reader.metadata = {}
        mock_reader.pages = [mock_page]

        with patch("pypdf.PdfReader", return_value=mock_reader):
            _, meta = load_pdf(fake_pdf)

        assert meta.author is None
        assert meta.title is None

    def test_load_docx_extracts_author_and_title(self, tmp_path):
        """load_docx reads author and title from core_properties."""
        try:
            import docx as docx_lib  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        import docx as docx_lib

        doc_file = tmp_path / "book.docx"
        wdoc = docx_lib.Document()
        wdoc.core_properties.author = "Leo Tolstoy"
        wdoc.core_properties.title = "War and Peace"
        wdoc.add_paragraph("Chapter 1")
        wdoc.add_paragraph("The story begins on a dark night.")
        wdoc.save(str(doc_file))

        from storysphere.pipelines.document_processing.loader import load_docx

        segments, meta = load_docx(doc_file)

        assert meta.author == "Leo Tolstoy"
        assert meta.title == "War and Peace"
        assert len(segments) > 0

    def test_load_docx_handles_empty_metadata(self, tmp_path):
        """load_docx returns None fields when core_properties are blank."""
        try:
            import docx as docx_lib  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        import docx as docx_lib

        doc_file = tmp_path / "blank_meta.docx"
        wdoc = docx_lib.Document()
        wdoc.add_paragraph("Some content here.")
        wdoc.save(str(doc_file))

        from storysphere.pipelines.document_processing.loader import load_docx

        _, meta = load_docx(doc_file)

        # python-docx defaults: author may be system username, title is empty
        assert meta.title is None

    def test_load_txt_detects_big5_encoding(self, tmp_path):
        """load_txt auto-detects Big5-encoded Traditional Chinese text."""
        from storysphere.pipelines.document_processing.loader import load_txt

        txt_file = tmp_path / "novel_big5.txt"
        txt_file.write_bytes("第一章 你好，世界！這是繁體中文測試。".encode("big5"))

        segments, _ = load_txt(txt_file)

        assert len(segments) > 0
        assert "你好" in segments[0][1]

    def test_load_txt_detects_gbk_encoding(self, tmp_path):
        """load_txt auto-detects GBK-encoded Simplified Chinese text."""
        from storysphere.pipelines.document_processing.loader import load_txt

        txt_file = tmp_path / "novel_gbk.txt"
        txt_file.write_bytes("第一章 你好，世界！这是简体中文测试。".encode("gbk"))

        segments, _ = load_txt(txt_file)

        assert len(segments) > 0
        assert "你好" in segments[0][1]

    def test_load_txt_reads_utf8_as_before(self, tmp_path):
        """load_txt still reads plain UTF-8 text correctly without an explicit encoding."""
        from storysphere.pipelines.document_processing.loader import load_txt

        txt_file = tmp_path / "novel_utf8.txt"
        txt_file.write_text("Chapter 1\nIt was a dark and stormy night.", encoding="utf-8")

        segments, _ = load_txt(txt_file)

        assert segments[0][1] == "Chapter 1"
        assert segments[1][1] == "It was a dark and stormy night."

    def test_load_txt_explicit_encoding_still_supported(self, tmp_path):
        """Passing an explicit encoding bypasses auto-detection, matching prior behavior."""
        from storysphere.pipelines.document_processing.loader import load_txt

        txt_file = tmp_path / "novel_explicit.txt"
        txt_file.write_text("Hello", encoding="utf-16")

        segments, _ = load_txt(txt_file, encoding="utf-16")

        assert segments[0][1] == "Hello"

    def test_load_docx_records_heading_style_indices(self, tmp_path):
        """load_docx flags paragraph indices using a built-in Heading style."""
        try:
            import docx as docx_lib
        except ImportError:
            pytest.skip("python-docx not installed")

        doc_file = tmp_path / "styled.docx"
        wdoc = docx_lib.Document()
        wdoc.add_heading("楔子", level=1)
        wdoc.add_paragraph("很久很久以前。")
        wdoc.add_heading("第一章 開始", level=1)
        wdoc.add_paragraph("故事開始了。")
        wdoc.save(str(doc_file))

        from storysphere.pipelines.document_processing.loader import load_docx

        segments, meta = load_docx(doc_file)

        heading_texts = {text for idx, text in segments if idx in meta.heading_indices}
        assert heading_texts == {"楔子", "第一章 開始"}

    def test_load_docx_no_headings_gives_empty_indices(self, tmp_path):
        """A DOCX with only "Normal"-style paragraphs flags no heading indices."""
        try:
            import docx as docx_lib
        except ImportError:
            pytest.skip("python-docx not installed")

        doc_file = tmp_path / "plain.docx"
        wdoc = docx_lib.Document()
        wdoc.add_paragraph("Just a plain paragraph.")
        wdoc.save(str(doc_file))

        from storysphere.pipelines.document_processing.loader import load_docx

        _, meta = load_docx(doc_file)

        assert meta.heading_indices == set()

    def test_load_pdf_strips_repeated_running_header(self, tmp_path):
        """A line repeated at the top of most pages is filtered as a running header."""
        from storysphere.pipelines.document_processing.loader import load_pdf

        fake_pdf = tmp_path / "novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        page_texts = [
            "書名：測試小說\n正文第一段內容。\n1",
            "書名：測試小說\n正文第二段內容。\n2",
            "書名：測試小說\n正文第三段內容。\n3",
            "書名：測試小說\n正文第四段內容。\n4",
        ]
        mock_pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            mock_pages.append(page)

        mock_reader = MagicMock()
        mock_reader.metadata = {}
        mock_reader.pages = mock_pages

        with patch("pypdf.PdfReader", return_value=mock_reader):
            segments, _ = load_pdf(fake_pdf)

        all_text = [text for _, text in segments]
        assert "書名：測試小說" not in all_text
        assert "正文第一段內容。" in all_text
        assert "正文第四段內容。" in all_text

    def test_load_pdf_keeps_boundary_line_seen_on_few_pages(self, tmp_path):
        """A boundary line that only repeats on a couple of pages is kept —
        the recurrence threshold avoids false positives."""
        from storysphere.pipelines.document_processing.loader import load_pdf

        fake_pdf = tmp_path / "novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        page_texts = [
            "特殊標語\n正文第一段內容。",
            "正文第二段內容。\n特殊標語",
            "正文第三段內容。\n正文第四段內容。",
        ]
        mock_pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            mock_pages.append(page)

        mock_reader = MagicMock()
        mock_reader.metadata = {}
        mock_reader.pages = mock_pages

        with patch("pypdf.PdfReader", return_value=mock_reader):
            segments, _ = load_pdf(fake_pdf)

        all_text = [text for _, text in segments]
        assert "特殊標語" in all_text

    def test_load_pdf_few_pages_skips_header_footer_filtering(self, tmp_path):
        """Documents with under 3 pages never trigger the noise filter,
        matching the existing single-page metadata tests' expectations."""
        from storysphere.pipelines.document_processing.loader import load_pdf

        fake_pdf = tmp_path / "novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        page_texts = ["重複行\n正文一。", "重複行\n正文二。"]
        mock_pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            mock_pages.append(page)

        mock_reader = MagicMock()
        mock_reader.metadata = {}
        mock_reader.pages = mock_pages

        with patch("pypdf.PdfReader", return_value=mock_reader):
            segments, _ = load_pdf(fake_pdf)

        all_text = [text for _, text in segments]
        assert all_text.count("重複行") == 2


# ── DocumentProcessingPipeline: metadata propagation ────────────────────────


class TestDocumentProcessingPipelineMeta:
    @pytest.mark.asyncio
    async def test_metadata_title_takes_priority_over_filename(self, tmp_path):
        """When PDF metadata has a title, it overrides the filename stem."""
        from storysphere.pipelines.document_processing.loader import DocumentMeta
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        fake_pdf = tmp_path / "untitled_file.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        meta = DocumentMeta(title="Great Expectations", author="Charles Dickens")
        segments = [(0, "Chapter 1"), (1, "It was the best of times.")]

        with patch(
            "storysphere.pipelines.document_processing.pipeline.DocumentProcessingPipeline._load_sync",
            return_value=(segments, meta),
        ):
            pipeline = DocumentProcessingPipeline()
            result = await pipeline.run(fake_pdf)

        assert result.title == "Great Expectations"
        assert result.author == "Charles Dickens"

    @pytest.mark.asyncio
    async def test_filename_stem_used_when_metadata_title_absent(self, tmp_path):
        """When metadata has no title, filename stem is used."""
        from storysphere.pipelines.document_processing.loader import DocumentMeta
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        fake_pdf = tmp_path / "my_novel.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        meta = DocumentMeta(title=None, author=None)
        segments = [(0, "Chapter 1"), (1, "Once upon a time in a land far away.")]

        with patch(
            "storysphere.pipelines.document_processing.pipeline.DocumentProcessingPipeline._load_sync",
            return_value=(segments, meta),
        ):
            pipeline = DocumentProcessingPipeline()
            result = await pipeline.run(fake_pdf)

        assert result.title == "my_novel"
        assert result.author is None

    @pytest.mark.asyncio
    async def test_author_is_none_when_metadata_absent(self, tmp_path):
        """doc.author stays None when loader finds no author in metadata."""
        from storysphere.pipelines.document_processing.loader import DocumentMeta
        from storysphere.pipelines.document_processing.pipeline import DocumentProcessingPipeline

        fake_pdf = tmp_path / "anonymous.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4 stub")

        meta = DocumentMeta(title="Some Title", author=None)
        segments = [(0, "Chapter 1"), (1, "The protagonist walked slowly down the hall.")]

        with patch(
            "storysphere.pipelines.document_processing.pipeline.DocumentProcessingPipeline._load_sync",
            return_value=(segments, meta),
        ):
            pipeline = DocumentProcessingPipeline()
            result = await pipeline.run(fake_pdf)

        assert result.author is None

